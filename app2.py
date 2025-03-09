import openai
import streamlit as st
import streamlit.components.v1 as components
import pdfplumber
import docx
from pydantic import BaseModel
from typing import List
import re
import markdown
from io import BytesIO
from docx import Document
from xhtml2pdf import pisa

# Streamlit UI
st.title("AI Compliance Training Script Generator")

# User Input for OpenAI API Key
st.subheader("OpenAI API Key")
openai_api_key = st.text_input("Enter your OpenAI API Key:", type="password", placeholder="sk-...", key="api_key")

if not openai_api_key:
    st.warning("Please enter your OpenAI API key to generate scripts.")
    st.stop()

# Set OpenAI API Key
openai.api_key = openai_api_key

# File Upload Section
st.subheader("Upload Course Outline (PDF, DOCX, TXT)")
uploaded_files = st.file_uploader("Upload multiple course outlines", type=["pdf", "docx", "txt"], accept_multiple_files=True)

# User-defined module identifier
st.subheader("Module Identifier")
module_identifier = st.text_input("Enter module identifier (e.g., 'Module', 'Section', 'Topic'):", value="Module")

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    try:
        with pdfplumber.open(pdf_file) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    except Exception as e:
        return f"Error extracting text from PDF: {e}"

# Define Pydantic models
class ModuleDetail(BaseModel):
    title: str
    content: str
    custom_prompt: str

class CourseDetail(BaseModel):
    title: str
    description: str
    duration_minutes: int
    audience: str
    regulations: str
    modules: List[ModuleDetail]

# Function to extract modules
def extract_modules_from_text(text, identifier):
    modules = []
    pattern = rf"\b{re.escape(identifier)}\s+\d+[:.-]?"
    module_matches = re.split(pattern, text, flags=re.IGNORECASE)

    if len(module_matches) > 1:
        for idx, section in enumerate(module_matches[1:], start=1):
            lines = section.strip().split("\n")
            module_title = lines[0].strip() if lines else f"{identifier} {idx} (Untitled)"
            module_content = "\n".join(lines[1:]).strip()
            modules.append(ModuleDetail(title=module_title, content=module_content, custom_prompt=""))
    return modules

# Extract content from uploaded documents
course_modules = []
if uploaded_files:
    for file in uploaded_files:
        file_type = file.name.split(".")[-1]
        extracted_text = extract_text_from_pdf(file) if file_type == "pdf" else extract_text_from_docx(file) if file_type == "docx" else file.read().decode("utf-8")

        if not extracted_text.strip():
            st.error(f"Could not extract text from {file.name}. Please check the file format.")
            continue

        st.text_area(f"Extracted Course Outline from {file.name}", extracted_text, height=300)
        detected_modules = extract_modules_from_text(extracted_text, module_identifier)

        if detected_modules:
            course_modules.extend(detected_modules)
        else:
            st.warning(f"No modules detected in {file.name} using '{module_identifier}'.")

# Course Metadata Inputs
st.subheader("Course Metadata")
title = st.text_input("Course Title:", placeholder="e.g., 2025 Compliance Essentials for RIAs – Annual Training")
description = st.text_area("Course Description:", placeholder="Provide a brief course summary.")
duration = st.number_input("Course Duration (minutes):", min_value=1, step=1, value=30)
audience = st.selectbox("Intended Audience:", ["RIA Employees", "Compliance Officers", "Investment Advisers", "General Finance Professionals"])
regulations = st.selectbox("Regulatory Alignment:", ["SEC", "FINRA", "Investment Advisers Act of 1940", "Multiple"])

# Function to generate module script
def generate_module_script(module: ModuleDetail, module_number: int) -> str:
    try:
        base_prompt = f"""
        Expand {module_identifier} {module_number} on "{module.title}" in compliance training.
        The module should be 700-1,000 words and include:

        **Learning Objectives:**
        - Define the compliance issue.
        - Provide practical applications.

        **Content:**
        - Overview of the compliance issue.
        - Key regulatory requirements and best practices.
        - Case study illustrating real-world application.

        **Scenario-Based Learning Activity:**
        - Provide a scenario related to "{module.title}".

        **Regulatory References:**
        - Cite SEC, FINRA, or applicable regulations.

        **Duration:**
        - This module should take approximately {duration} minutes to complete.

        {module.content}
        """

        # Allow users to override prompt
        prompt = module.custom_prompt if module.custom_prompt else base_prompt

        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "You are an expert in compliance training script generation."},
                      {"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.2
        )

        return response["choices"][0]["message"]["content"].strip()

    except Exception as e:
        return f"Error generating {module_identifier} {module_number}: {e}"

# Function to convert HTML to DOCX and return binary data
def html_to_docx(html_content):
    doc = Document()
    doc.add_heading('Compliance Training Script', 0)
    doc.add_paragraph(html_content, style='BodyText')
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()

# Function to convert HTML to PDF
def html_to_pdf(html_content):
    pdf_buffer = BytesIO()
    pisa.CreatePDF(BytesIO(html_content.encode('utf-8')), dest=pdf_buffer)
    return pdf_buffer

# CKEditor Function with content capture
def ckeditor(md_text):
    html_content = markdown.markdown(md_text)  # Convert Markdown to HTML
    ckeditor_html = f"""
    <head>
        <script src="https://cdn.ckeditor.com/4.22.1/full/ckeditor.js"></script>
    </head>
    <body>
        <textarea id="editor">{html_content}</textarea>
        <script>
            CKEDITOR.replace('editor', {{
                extraPlugins: 'exportpdf',
                toolbar: [
                    ['Save', 'Print', '-', 'Undo', 'Redo'],
                    ['Bold', 'Italic', 'Underline', 'Strike'],
                    ['NumberedList', 'BulletedList'],
                    ['Link', 'Unlink'],
                    ['ExportPdf', 'Source']
                ],
                removeButtons: ''
            }});

            // Function to handle Save
            CKEDITOR.instances.editor.on('save', function(evt) {{
                alert('Content saved inside CKEditor. Use export options to download.');
            }});
        </script>
    </body>
    """
    components.html(ckeditor_html, height=500)

# Function to generate detailed storyboard
def generate_detailed_storyboard(module: ModuleDetail, module_number: int) -> str:
    storyboard_prompt = f"""
    Create a detailed storyboard for an eLearning course based on the provided course outline. The storyboard should include slide-by-slide visuals and descriptions that align with each module's learning objectives.

    For each module and topic, include the following:

    1. Slide Number and Title – A concise title summarizing the slide content.
    2. Visual Description – Describe the imagery, animations, or graphics that should appear.
    3. Text Overlay – Key messages or text elements displayed on the slide.
    4. Interactive Elements (if applicable) – Any animations, checklists, flowcharts, quizzes, or case studies.

    Ensure the storyboard is engaging, visually structured, and aligns with compliance training best practices. Include case study highlights, real-world examples, infographics, and animations where relevant.

    Module {module_number}: {module.title}
    {module.content}
    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "You are an expert in eLearning storyboard creation."},
                      {"role": "user", "content": storyboard_prompt}],
            max_tokens=2000,
            temperature=0.2
        )

        storyboard = response["choices"][0]["message"]["content"].strip()
        return f"Module {module_number}: {module.title}\n\n{storyboard}"

    except Exception as e:
        return f"Error generating detailed storyboard for {module_identifier} {module_number}: {e}"

# Generate Training Script Button
if st.button("Generate Training Script"):
    if not title or not description or not audience or not regulations or not course_modules:
        st.warning("Please fill in all fields before generating the script!")
    else:
        full_script = ""
        full_storyboard = ""
        for idx, module in enumerate(course_modules, start=1):
            with st.spinner(f"Generating {module_identifier} {idx}: {module.title}..."):
                module_script = generate_module_script(module, idx)
                module_storyboard = generate_detailed_storyboard(module, idx)
                full_script += module_script + "\n\n"
                full_storyboard += module_storyboard + "\n\n"

        st.subheader("Edit Your Script Before Exporting")
        ckeditor(full_script)  # Convert Markdown to HTML

        st.subheader("Detailed Storyboard")
        st.text_area("Generated Detailed Storyboard", full_storyboard, height=300)
